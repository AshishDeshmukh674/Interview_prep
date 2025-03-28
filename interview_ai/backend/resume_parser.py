import logging
import io
import pdfplumber
from docx import Document
from typing import Dict, Any, List
import re
from pathlib import Path
import asyncio

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ResumeParser:
    def __init__(self):
        self.supported_extensions = {'.pdf', '.docx', '.txt'}

    async def parse_resume(self, file_path: Path) -> Dict[str, Any]:
        """
        Parse resume from various file formats.
        """
        try:
            # Validate file extension
            if file_path.suffix.lower() not in self.supported_extensions:
                raise ValueError(f"Unsupported file format: {file_path.suffix}")
            
            # Extract text based on file type
            text = await self._extract_text(file_path)
            
            # Parse the extracted text
            resume_data = self._parse_text(text)
            
            return resume_data
            
        except Exception as e:
            logger.error(f"Error parsing resume: {str(e)}")
            raise

    async def _extract_text(self, file_path: Path) -> str:
        """
        Extract text from resume file based on its format.
        """
        try:
            if file_path.suffix.lower() == '.pdf':
                return await self._extract_from_pdf(file_path)
            elif file_path.suffix.lower() == '.docx':
                return await self._extract_from_docx(file_path)
            elif file_path.suffix.lower() == '.txt':
                return await self._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_path.suffix}")
                
        except Exception as e:
            logger.error(f"Error extracting text: {str(e)}")
            raise

    async def _extract_from_pdf(self, file_path: Path) -> str:
        """
        Extract text from PDF file.
        """
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise

    async def _extract_from_docx(self, file_path: Path) -> str:
        """
        Extract text from DOCX file.
        """
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise

    async def _extract_from_txt(self, file_path: Path) -> str:
        """
        Extract text from TXT file.
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
                
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {str(e)}")
            raise

    def _parse_text(self, text: str) -> Dict[str, Any]:
        """
        Parse extracted text into structured resume data.
        """
        try:
            # Split text into sections
            sections = self._split_into_sections(text)
            
            # Extract information from each section
            resume_data = {
                "personal_info": self._extract_personal_info(sections.get("personal", "")),
                "education": self._extract_education(sections.get("education", "")),
                "experience": self._extract_experience(sections.get("experience", "")),
                "skills": self._extract_skills(sections.get("skills", "")),
                "projects": self._extract_projects(sections.get("projects", ""))
            }
            
            return resume_data
            
        except Exception as e:
            logger.error(f"Error parsing text: {str(e)}")
            raise

    def _split_into_sections(self, text: str) -> Dict[str, str]:
        """
        Split resume text into sections.
        """
        sections = {}
        current_section = "personal"
        current_text = []
        
        # Common section headers
        section_headers = {
            "education": ["education", "academic", "qualification"],
            "experience": ["experience", "work", "employment", "career"],
            "skills": ["skills", "expertise", "technologies", "technical"],
            "projects": ["projects", "portfolio", "work samples"]
        }
        
        # Split text into lines
        lines = text.split("\n")
        
        for line in lines:
            line = line.strip().lower()
            
            # Check if line is a section header
            is_header = False
            for section, keywords in section_headers.items():
                if any(keyword in line for keyword in keywords):
                    sections[current_section] = "\n".join(current_text)
                    current_section = section
                    current_text = []
                    is_header = True
                    break
            
            if not is_header:
                current_text.append(line)
        
        # Add the last section
        sections[current_section] = "\n".join(current_text)
        
        return sections

    def _extract_personal_info(self, text: str) -> Dict[str, str]:
        """
        Extract personal information from text.
        """
        info = {
            "name": "",
            "email": "",
            "phone": "",
            "location": "",
            "linkedin": ""
        }
        
        # Extract email
        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        emails = re.findall(email_pattern, text)
        if emails:
            info["email"] = emails[0]
        
        # Extract phone
        phone_pattern = r'\+?[\d\s-()]{10,}'
        phones = re.findall(phone_pattern, text)
        if phones:
            info["phone"] = phones[0]
        
        # Extract LinkedIn URL
        linkedin_pattern = r'linkedin\.com/in/[a-zA-Z0-9-]+'
        linkedin_urls = re.findall(linkedin_pattern, text)
        if linkedin_urls:
            info["linkedin"] = linkedin_urls[0]
        
        # First line is usually the name
        lines = text.split("\n")
        if lines:
            info["name"] = lines[0].strip()
        
        return info

    def _extract_education(self, text: str) -> List[Dict[str, str]]:
        """
        Extract education information from text.
        """
        education = []
        current_entry = {}
        
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                if current_entry:
                    education.append(current_entry)
                    current_entry = {}
                continue
            
            # Try to identify degree and institution
            if not current_entry.get("degree"):
                current_entry["degree"] = line
            elif not current_entry.get("institution"):
                current_entry["institution"] = line
        
        if current_entry:
            education.append(current_entry)
        
        return education

    def _extract_experience(self, text: str) -> List[Dict[str, str]]:
        """
        Extract work experience from text.
        """
        experience = []
        current_entry = {}
        
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                if current_entry:
                    experience.append(current_entry)
                    current_entry = {}
                continue
            
            # Try to identify company and role
            if not current_entry.get("company"):
                current_entry["company"] = line
            elif not current_entry.get("role"):
                current_entry["role"] = line
            else:
                # Add to description
                if "description" not in current_entry:
                    current_entry["description"] = []
                current_entry["description"].append(line)
        
        if current_entry:
            experience.append(current_entry)
        
        return experience

    def _extract_skills(self, text: str) -> List[str]:
        """
        Extract skills from text.
        """
        skills = []
        
        # Split by common delimiters
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue
            
            # Split by commas, semicolons, or spaces
            line_skills = [s.strip() for s in line.replace(",", " ").replace(";", " ").split()]
            skills.extend(line_skills)
        
        # Remove duplicates and empty strings
        return list(set(filter(None, skills)))

    def _extract_projects(self, text: str) -> List[Dict[str, str]]:
        """
        Extract project information from text.
        """
        projects = []
        current_project = {}
        
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                if current_project:
                    projects.append(current_project)
                    current_project = {}
                continue
            
            # Try to identify project name and description
            if not current_project.get("name"):
                current_project["name"] = line
            else:
                # Add to description
                if "description" not in current_project:
                    current_project["description"] = []
                current_project["description"].append(line)
        
        if current_project:
            projects.append(current_project)
        
        return projects

def extract_resume_data(file_content: bytes) -> Dict[str, Any]:
    """
    Extract information from a resume file.
    """
    try:
        # Create a BytesIO object from the file content
        file_stream = io.BytesIO(file_content)
        
        # Try to read as PDF first
        try:
            with pdfplumber.open(file_stream) as pdf:
                text = ""
                for page in pdf.pages:
                    text += page.extract_text() or ""
        except Exception as e:
            logger.warning(f"Failed to read as PDF: {str(e)}")
            # Try to read as DOCX
            try:
                file_stream.seek(0)  # Reset stream position
                doc = Document(file_stream)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            except Exception as e:
                logger.error(f"Failed to read as DOCX: {str(e)}")
                raise ValueError("Unsupported file format. Please upload a PDF or DOCX file.")
        
        # Extract information from text
        return {
            "resume_text": text,
            "contact_info": _extract_contact_info(text),
            "education": _extract_education(text),
            "experience": _extract_experience(text),
            "skills": _extract_skills(text)
        }
        
    except Exception as e:
        logger.error(f"Error extracting resume data: {str(e)}")
        raise

def _extract_contact_info(text: str) -> Dict[str, str]:
    """
    Extract contact information from resume text.
    """
    contact_info = {
        "name": "",
        "email": "",
        "phone": "",
        "location": ""
    }
    
    # Extract email
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    email_match = re.search(email_pattern, text)
    if email_match:
        contact_info["email"] = email_match.group()
    
    # Extract phone
    phone_pattern = r'(?:\+\d{1,3}[-. ]?)?\(?\d{3}\)?[-. ]?\d{3}[-. ]?\d{4}'
    phone_match = re.search(phone_pattern, text)
    if phone_match:
        contact_info["phone"] = phone_match.group()
    
    # Extract name (first line of text)
    lines = text.split('\n')
    if lines:
        contact_info["name"] = lines[0].strip()
    
    # Extract location (look for common location indicators)
    location_pattern = r'(?:[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*,\s*[A-Z]{2})'
    location_match = re.search(location_pattern, text)
    if location_match:
        contact_info["location"] = location_match.group()
    
    return contact_info

def _extract_education(text: str) -> List[Dict[str, str]]:
    """
    Extract education information from resume text.
    """
    education = []
    
    # Look for education section
    education_section = _find_section(text, ["Education", "Academic Background", "Qualifications"])
    if not education_section:
        return education
    
    # Split into individual entries
    entries = education_section.split('\n\n')
    
    for entry in entries:
        if not entry.strip():
            continue
            
        # Try to extract degree and institution
        degree_match = re.search(r'([A-Za-z\s]+(?:Bachelor|Master|Doctor|PhD|B\.?S\.?|M\.?S\.?|Ph\.?D\.?)[A-Za-z\s]+)', entry)
        if degree_match:
            education.append({
                "degree": degree_match.group(1).strip(),
                "institution": entry.split('\n')[0].strip(),
                "year": _extract_year(entry)
            })
    
    return education

def _extract_experience(text: str) -> List[Dict[str, str]]:
    """
    Extract work experience from resume text.
    """
    experience = []
    
    # Look for experience section
    experience_section = _find_section(text, ["Experience", "Work Experience", "Professional Experience"])
    if not experience_section:
        return experience
    
    # Split into individual entries
    entries = experience_section.split('\n\n')
    
    for entry in entries:
        if not entry.strip():
            continue
            
        # Try to extract company and position
        lines = entry.split('\n')
        if len(lines) >= 2:
            experience.append({
                "position": lines[0].strip(),
                "company": lines[1].strip(),
                "duration": _extract_duration(entry),
                "description": '\n'.join(lines[2:]).strip()
            })
    
    return experience

def _extract_skills(text: str) -> List[str]:
    """
    Extract skills from resume text.
    """
    skills = []
    
    # Look for skills section
    skills_section = _find_section(text, ["Skills", "Technical Skills", "Core Competencies"])
    if not skills_section:
        return skills
    
    # Split into individual skills
    skill_entries = skills_section.split('\n')
    
    for entry in skill_entries:
        # Split by common delimiters
        entry_skills = re.split(r'[,•]', entry)
        for skill in entry_skills:
            skill = skill.strip()
            if skill and len(skill) > 1:  # Avoid empty or single-character skills
                skills.append(skill)
    
    return list(set(skills))  # Remove duplicates

def _find_section(text: str, section_names: List[str]) -> str:
    """
    Find a section in the resume text by looking for section headers.
    """
    lines = text.split('\n')
    section_start = -1
    section_end = -1
    
    for i, line in enumerate(lines):
        line = line.strip().lower()
        if any(name.lower() in line for name in section_names):
            section_start = i
            break
    
    if section_start == -1:
        return ""
    
    # Look for next section
    for i in range(section_start + 1, len(lines)):
        line = lines[i].strip()
        if line and line.isupper():  # Common section header format
            section_end = i
            break
    
    if section_end == -1:
        section_end = len(lines)
    
    return '\n'.join(lines[section_start + 1:section_end])

def _extract_year(text: str) -> str:
    """
    Extract year from text.
    """
    year_match = re.search(r'\b(19|20)\d{2}\b', text)
    return year_match.group() if year_match else ""

def _extract_duration(text: str) -> str:
    """
    Extract duration from text.
    """
    duration_pattern = r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}\s*-\s*(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}|\d{4}\s*-\s*(?:present|current)'
    duration_match = re.search(duration_pattern, text, re.IGNORECASE)
    return duration_match.group() if duration_match else ""
